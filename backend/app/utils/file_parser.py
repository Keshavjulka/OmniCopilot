"""
File parsing utilities for PDF and DOCX extraction.
"""

import io
import logging
from typing import Union

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_obj: Union[io.BytesIO, bytes]) -> str:
    """Extract text from a PDF file."""
    try:
        import pdfplumber
        if isinstance(file_obj, bytes):
            file_obj = io.BytesIO(file_obj)
        text_parts = []
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts)
    except ImportError:
        # Fallback to PyPDF2
        try:
            import PyPDF2
            if isinstance(file_obj, bytes):
                file_obj = io.BytesIO(file_obj)
            reader = PyPDF2.PdfReader(file_obj)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            logger.error("No PDF library found. Install pdfplumber or PyPDF2.")
            return "[PDF text extraction unavailable]"


def extract_text_from_docx(file_obj: Union[io.BytesIO, bytes]) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        if isinstance(file_obj, bytes):
            file_obj = io.BytesIO(file_obj)
        doc = Document(file_obj)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except ImportError:
        logger.error("python-docx not installed.")
        return "[DOCX text extraction unavailable]"
